#!/usr/bin/env python3
"""Markdown -> Hebrew RTL .docx renderer for Pnina's tailored CVs.

Schema (each section separated by a `---` line on its own):

  Block 1: header — `**<name>**` + contact line
  Block 2: `**תקציר מקצועי**` + paragraph
  Block 3: `**ניסיון מקצועי**` + role headers `**T** | **O  |  D**` + `* bullets`
  Block 4: `**השכלה**` + lines `**degree** – **institution**`
  Block 5: `**כישורים מרכזיים**` + `* bullets`
  Block 6: `**שפות ורישיונות**` + free lines
  Block 7: footer line starting with `*`

Failure contract: on any error, print a single Hebrew line to stderr and exit 1.
Never emit a Python traceback. Callers (the agent) pass stderr verbatim to WhatsApp.
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = 'Calibri'
NAME_COLOR = (0x1F, 0x38, 0x96)
SECTION_COLOR = (0x2E, 0x74, 0xB5)
SECTION_COLOR_HEX = '2E74B5'
CONTACT_COLOR = (0x40, 0x40, 0x40)
FOOTER_COLOR = (0x80, 0x80, 0x80)

SECTION_HEADINGS = {
    'תקציר מקצועי',
    'ניסיון מקצועי',
    'השכלה',
    'כישורים מרכזיים',
    'שפות ורישיונות',
}

REQUIRED_SECTIONS = [
    'תקציר מקצועי',
    'ניסיון מקצועי',
    'השכלה',
    'כישורים מרכזיים',
    'שפות ורישיונות',
]


class RenderError(Exception):
    """Caught at the top level; message is printed verbatim to stderr in Hebrew."""


# ---------------------------------------------------------------------------
# Helper functions — lifted verbatim (with minor parameterization) from
# resumes/generate_shkulo_tov_shikum_deputy_docx.py. Do not change the rendering
# decisions here without re-validating against the canonical .docx.
# ---------------------------------------------------------------------------

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)


def add_rtl_paragraph(doc, text, bold=False, font_size=11, space_before=0,
                      space_after=4, color=None):
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = FONT
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    return para


def add_bullet(doc, text, font_size=11):
    para = doc.add_paragraph(style='List Bullet')
    set_rtl(para)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:right'), '360')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = FONT
    set_rtl_run(run)
    return para


def add_divider(doc, color=SECTION_COLOR_HEX):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_title(doc, text):
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT
    run.font.color.rgb = RGBColor(*SECTION_COLOR)
    set_rtl_run(run)
    add_divider(doc)
    return para


def add_role_header(doc, title, org_dates):
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    r1 = para.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = FONT
    set_rtl_run(r1)
    r2 = para.add_run('  |  ' + org_dates)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.name = FONT
    r2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    set_rtl_run(r2)
    return para


def add_education_entry(doc, degree, inst):
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    r1 = para.add_run(degree)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = FONT
    set_rtl_run(r1)
    r2 = para.add_run(' – ')
    r2.font.size = Pt(11)
    r2.font.name = FONT
    set_rtl_run(r2)
    r3 = para.add_run(inst)
    r3.bold = True
    r3.font.size = Pt(11)
    r3.font.name = FONT
    set_rtl_run(r3)
    return para


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

_BOLD = r'\*\*(.+?)\*\*'
_ROLE_HEADER_RE = re.compile(rf'^{_BOLD}\s+\|\s+{_BOLD}\s*$')
_EDUCATION_LINE_RE = re.compile(rf'^{_BOLD}\s+–\s+{_BOLD}\s*$')


def strip_bold(s: str) -> str:
    s = s.strip()
    m = re.fullmatch(_BOLD, s)
    return m.group(1) if m else s


def parse_role_header(line: str):
    """Return (title, org_dates) from a role-header line.

    Strict schema: `**<title>** | **<org>  |  <dates>**`.
    Raises RenderError with a Hebrew message on any deviation.
    """
    m = _ROLE_HEADER_RE.match(line.strip())
    if not m:
        raise RenderError(
            f"שגיאה בהפקת קובץ: שורת תפקיד לא תקינה — חסר ' | ' בין כותרת לארגון: {line.strip()}"
        )
    title, org_dates = m.group(1), m.group(2)
    if '  |  ' not in org_dates:
        raise RenderError(
            f"שגיאה בהפקת קובץ: שורת תפקיד לא תקינה — חסר '  |  ' בין הארגון לתאריכים: {line.strip()}"
        )
    return title, org_dates


def parse_education_line(line: str):
    """Return (degree, institution) from `**degree** – **institution**`."""
    m = _EDUCATION_LINE_RE.match(line.strip())
    if not m:
        raise RenderError(
            f"שגיאה בהפקת קובץ: שורת השכלה לא תקינה — חסר ' – ' בין תואר למוסד: {line.strip()}"
        )
    return m.group(1), m.group(2)


def parse_sections(text: str) -> dict:
    """Split the markdown on `---` lines and bucket each block by section heading.

    Returns:
      {
        'header_block': str,            # raw text of the header block
        'תקציר מקצועי': str,            # section body
        'ניסיון מקצועי': str,
        ...
        'footer_block': str,            # the trailing `*...` line
      }
    """
    blocks = []
    current = []
    for line in text.splitlines():
        if line.strip() == '---':
            blocks.append('\n'.join(current).strip())
            current = []
        else:
            current.append(line)
    if current:
        blocks.append('\n'.join(current).strip())

    if not blocks:
        raise RenderError("שגיאה בהפקת קובץ: המסמך ריק")

    out = {'header_block': blocks[0]}

    for block in blocks[1:]:
        lines = [l for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        heading = strip_bold(lines[0])
        if heading in SECTION_HEADINGS:
            body = '\n'.join(lines[1:]).strip()
            out[heading] = body
        elif lines[0].lstrip().startswith('*') and 'footer_block' not in out:
            out['footer_block'] = lines[0].strip()
        # else: unknown block — silently ignored. The agent should not emit these.

    for required in REQUIRED_SECTIONS:
        if required not in out:
            raise RenderError(f"שגיאה בהפקת קובץ: סעיף '{required}' חסר במסמך המקור")

    return out


def parse_header(block: str):
    """Return (name, contact_line) from the first block."""
    lines = [l for l in block.splitlines() if l.strip()]
    if not lines:
        raise RenderError("שגיאה בהפקת קובץ: בלוק הכותרת ריק (שם איש קשר חסרים)")
    name = strip_bold(lines[0])
    if len(lines) < 2:
        raise RenderError("שגיאה בהפקת קובץ: שורת פרטי קשר חסרה")
    contact = lines[1].strip()
    return name, contact


def parse_bullets(body: str) -> list[str]:
    """Return list of bullet texts from lines starting with `*` or `* `."""
    out = []
    for line in body.splitlines():
        s = line.strip()
        if s.startswith('* '):
            out.append(s[2:].strip())
        elif s.startswith('*') and not s.startswith('**'):
            out.append(s[1:].strip())
    return out


def parse_experience(body: str) -> list[dict]:
    """Return list of {title, org_dates, bullets} for each role in the section."""
    roles = []
    current = None
    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if _ROLE_HEADER_RE.match(line):
            if current is not None:
                roles.append(current)
            title, org_dates = parse_role_header(line)
            current = {'title': title, 'org_dates': org_dates, 'bullets': []}
        elif line.startswith('* ') or (line.startswith('*') and not line.startswith('**')):
            if current is None:
                raise RenderError(
                    "שגיאה בהפקת קובץ: נמצאה תבליט לפני שורת תפקיד בסעיף 'ניסיון מקצועי'"
                )
            text = line[2:].strip() if line.startswith('* ') else line[1:].strip()
            current['bullets'].append(text)
        elif line.startswith('**') and '|' not in line:
            # Likely a malformed role header
            parse_role_header(line)  # will raise with a precise Hebrew message
        # else: silently skip unrecognized lines (blank etc.)
    if current is not None:
        roles.append(current)
    if not roles:
        raise RenderError("שגיאה בהפקת קובץ: לא נמצאו תפקידים בסעיף 'ניסיון מקצועי'")
    return roles


def parse_education(body: str) -> list[tuple]:
    """Return list of (degree, institution) pairs."""
    out = []
    for line in body.splitlines():
        s = line.strip()
        if not s:
            continue
        out.append(parse_education_line(s))
    if not out:
        raise RenderError("שגיאה בהפקת קובץ: לא נמצאו תארים בסעיף 'השכלה'")
    return out


def parse_languages(body: str) -> list[str]:
    """Return list of non-empty content lines for the languages-and-licenses section."""
    out = [l.strip() for l in body.splitlines() if l.strip()]
    if not out:
        raise RenderError("שגיאה בהפקת קובץ: סעיף 'שפות ורישיונות' ריק")
    return out


# ---------------------------------------------------------------------------
# Renderer
# ---------------------------------------------------------------------------

def render(input_path: Path, output_path: Path):
    text = input_path.read_text(encoding='utf-8')
    sections = parse_sections(text)
    name, contact = parse_header(sections['header_block'])
    summary = sections['תקציר מקצועי'].strip()
    if not summary:
        raise RenderError("שגיאה בהפקת קובץ: סעיף 'תקציר מקצועי' ריק")
    experience = parse_experience(sections['ניסיון מקצועי'])
    education = parse_education(sections['השכלה'])
    skills = parse_bullets(sections['כישורים מרכזיים'])
    if not skills:
        raise RenderError("שגיאה בהפקת קובץ: לא נמצאו פריטים בסעיף 'כישורים מרכזיים'")
    languages = parse_languages(sections['שפות ורישיונות'])
    footer = sections.get('footer_block', '').lstrip('*').strip()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # Name (centered, dark blue, bold, 18pt)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(4)
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = FONT
    name_run.font.color.rgb = RGBColor(*NAME_COLOR)

    # Contact line (centered, dark gray, 10pt)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(6)
    contact_run = contact_para.add_run(contact)
    contact_run.font.size = Pt(10)
    contact_run.font.name = FONT
    contact_run.font.color.rgb = RGBColor(*CONTACT_COLOR)

    # Summary
    add_section_title(doc, 'תקציר מקצועי')
    add_rtl_paragraph(doc, summary, font_size=11, space_after=4)

    # Experience
    add_section_title(doc, 'ניסיון מקצועי')
    for role in experience:
        add_role_header(doc, role['title'], role['org_dates'])
        for bullet in role['bullets']:
            add_bullet(doc, bullet)

    # Education
    add_section_title(doc, 'השכלה')
    for degree, inst in education:
        add_education_entry(doc, degree, inst)

    # Skills (two-column pairing)
    add_section_title(doc, 'כישורים מרכזיים')
    for i in range(0, len(skills), 2):
        pair = '• ' + skills[i]
        if i + 1 < len(skills):
            pair += '   |   • ' + skills[i + 1]
        add_rtl_paragraph(doc, pair, font_size=10, space_after=2)

    # Languages & Licenses (one line per content line)
    add_section_title(doc, 'שפות ורישיונות')
    for i, line in enumerate(languages):
        last = i == len(languages) - 1
        add_rtl_paragraph(doc, line, font_size=11, space_after=4 if last else 2)

    # Footer
    if footer:
        add_rtl_paragraph(doc, '*' + footer, font_size=9,
                          color=FOOTER_COLOR, space_after=0)

    doc.save(str(output_path))


# Path-argument allowlist. The agent assembles --input and --output from
# (LLM-transliterated) company and role names embedded in a Bash command.
# Even with quoted arguments, we want a hard guard at the renderer so a
# misrendered filename can never execute anything: reject any character
# outside [a-zA-Z0-9_./-].
_SAFE_PATH_RE = re.compile(r'^[a-zA-Z0-9_./-]+$')


def _validate_path_arg(label: str, value: str) -> None:
    if not _SAFE_PATH_RE.match(value):
        raise RenderError(
            f"שגיאה בהפקת קובץ: ארגומנט '{label}' מכיל תווים לא חוקיים — "
            f"מותרים רק [a-zA-Z0-9_./-]: {value}"
        )


def main():
    parser = argparse.ArgumentParser(description='Render a Hebrew CV markdown to .docx')
    parser.add_argument('--input', required=True, help='Path to the markdown CV')
    parser.add_argument('--output', required=True, help='Path for the output .docx')
    args = parser.parse_args()

    try:
        _validate_path_arg('--input', args.input)
        _validate_path_arg('--output', args.output)
        input_path = Path(args.input)
        if not input_path.exists():
            raise RenderError(f"שגיאה בהפקת קובץ: קובץ הקלט לא נמצא: {args.input}")
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        render(input_path, output_path)
        print(f"OK: {output_path}")
        return 0
    except RenderError as e:
        print(str(e), file=sys.stderr)
        return 1
    except Exception as e:
        # Any unexpected exception: surface a clean Hebrew message, not a traceback.
        print(f"שגיאה בהפקת קובץ: תקלה לא צפויה — {e.__class__.__name__}: {e}",
              file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
